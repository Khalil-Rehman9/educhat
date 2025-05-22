"""
DOCX document processor using python-docx for text extraction.
"""

import os
import sys
from datetime import datetime
from typing import Dict, Any, List, Optional
import docx

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from app.backend.document_processors.base_processor import BaseDocumentProcessor

class DOCXProcessor(BaseDocumentProcessor):
    def validate_doc_format(self):
        # Simple format check to ensure .docx extension
        if not self.file_path.lower().endswith('.docx'):
            raise ValueError("Invalid file extension for DOCXProcessor.")
    """
    Processor for Microsoft Word (DOCX) documents using python-docx.
    """
    
    def extract_text(self) -> str:
        # Additional validation for DOCX format
        self.validate_doc_format()
        """
        Extract text content from a DOCX document.
        
        Returns:
            The extracted text content as a string
        """
        try:
            text_content = ""
            
            # Open the DOCX file
            doc = docx.Document(self.file_path)
            
            # Extract text from paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Extract text from tables
            for i, table in enumerate(doc.tables):
                text_content += f"\n--- Table {i+1} ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
            
            return text_content.strip()
        
        except Exception as e:
            error_msg = f"Error extracting text from DOCX: {str(e)}"
            print(error_msg)
            return f"ERROR: {error_msg}"
    
    def extract_metadata(self) -> Dict[str, Any]:
        """
        Extract metadata from a DOCX document.
        
        Returns:
            Dictionary of metadata
        """
        metadata = {
            "file_path": self.file_path,
            "document_id": self.document_id,
            "processor": "DOCXProcessor",
            "extraction_date": datetime.now().isoformat(),
            "paragraph_count": 0,
            "table_count": 0,
            "title": None,
            "author": None,
            "comments": None,
            "keywords": None,
            "last_modified_by": None,
            "revision": None,
            "created": None,
            "modified": None
        }
        
        try:
            # Open the DOCX file
            doc = docx.Document(self.file_path)
            
            # Count elements
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            
            # Extract document properties
            core_properties = doc.core_properties
            
            if core_properties:
                if core_properties.title:
                    metadata["title"] = core_properties.title
                if core_properties.author:
                    metadata["author"] = core_properties.author
                if core_properties.comments:
                    metadata["comments"] = core_properties.comments
                if core_properties.keywords:
                    metadata["keywords"] = core_properties.keywords
                if core_properties.last_modified_by:
                    metadata["last_modified_by"] = core_properties.last_modified_by
                if core_properties.revision:
                    metadata["revision"] = core_properties.revision
                if core_properties.created:
                    metadata["created"] = core_properties.created.isoformat() if hasattr(core_properties.created, 'isoformat') else str(core_properties.created)
                if core_properties.modified:
                    metadata["modified"] = core_properties.modified.isoformat() if hasattr(core_properties.modified, 'isoformat') else str(core_properties.modified)
        
        except Exception as e:
            metadata["error"] = str(e)
        
        return metadata
    
    def extract_sections(self) -> List[Dict[str, Any]]:
        """
        Extract document sections with headings.
        
        Returns:
            List of dictionaries containing section data
        """
        sections = []
        
        try:
            doc = docx.Document(self.file_path)
            
            current_section = None
            section_content = ""
            
            for para in doc.paragraphs:
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    # Save previous section if it exists
                    if current_section:
                        sections.append({
                            "title": current_section,
                            "content": section_content.strip(),
                            "level": int(current_section_level)
                        })
                    
                    # Start new section
                    current_section = para.text
                    current_section_level = para.style.name.replace('Heading ', '')
                    section_content = ""
                else:
                    # Add content to current section
                    if current_section:
                        section_content += para.text + "\n"
            
            # Save the last section
            if current_section:
                sections.append({
                    "title": current_section,
                    "content": section_content.strip(),
                    "level": int(current_section_level) if current_section_level.isdigit() else 0
                })
        
        except Exception as e:
            print(f"Error extracting sections: {str(e)}")
        
        return sections 